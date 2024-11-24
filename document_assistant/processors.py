# processors.py
import re
from document_assistant.core import logger, MAX_FILE_SIZE
from pathlib import Path
from typing import Dict, List, Optional
import PyPDF2
import docx
import tempfile
import os
import subprocess
import streamlit as st

from document_assistant.models import SimilarityCalculator

class DocumentProcessor:
    """Main document processing class"""
    
    def __init__(self, model_manager=None):
        self.processors = {
            'pdf': PDFProcessor.process,
            'docx': DocxProcessor.process,
            'doc': DocProcessor.process,
            'txt': TxtProcessor.process
        }
        self.model_manager = model_manager

    def process_documents(self, files: List) -> None:
        """Process multiple documents"""
        try:
            if not self.model_manager or not self.model_manager.check_initialized():
                st.error("⚠️ Models not properly initialized")
                return

            for file in files:
                try:
                    if file.name not in st.session_state.documents:
                        self.process_single_document(file)
                except Exception as e:
                    logger.error(f"Error processing {file.name}: {str(e)}")
                    st.error(f"Failed to process {file.name}")
                    
        except Exception as e:
            logger.error(f"Document processing error: {str(e)}")
            st.error("An error occurred during document processing")

    def process_single_document(self, file) -> None:
        """Process a single document with all features"""
        try:
            progress = st.empty()
            progress.text(f"📄 Processing {file.name}...")
            
            # 1. Extract text and metadata
            doc_info = self._process_text(file)
            if not doc_info:
                return

            # 2. Extract citations and metadata
            if hasattr(self.model_manager, 'citation_extractor'):
                progress.text("📚 Extracting metadata...")
                metadata = self.model_manager.citation_extractor.extract_metadata(doc_info['content'])
                doc_info.update(metadata)

            # 3. Generate summary
            if hasattr(self.model_manager, 'summary_model'):
                progress.text("📝 Generating summary...")
                summary = self.model_manager.summary_model.generate_summary(doc_info['content'])
                if summary:
                    doc_info['summary'] = summary

            # 4. Generate visualization
            if hasattr(self.model_manager, 'image_generator'):
                progress.text("🎨 Creating visualization...")
                image, error = self.model_manager.image_generator.generate_image(
                    doc_info.get('summary', doc_info['content'][:500]),
                    doc_info.get('title', Path(file.name).stem)
                )
                if image:
                    doc_info['image'] = image
                elif error:
                    logger.warning(f"Image generation warning: {error}")

            # 5. Classify document
            if hasattr(self.model_manager, 'classifier'):
                progress.text("🏷️ Classifying content...")
                classification = self.model_manager.classifier.classify_document(doc_info['content'])
                if classification:
                    doc_info['classification'] = classification

            # Store processed document
            st.session_state.documents[file.name] = doc_info
            st.session_state.active_docs.add(file.name)

            # 6. Calculate similarities
            if len(st.session_state.documents) > 1 and hasattr(self.model_manager, 'similarity_model'):
                progress.text("🔄 Analyzing similarities...")
                self._update_similarities(file.name, doc_info)

            # Complete
            progress.empty()
            st.success(f"✅ Processed {file.name} successfully!")
            
        except Exception as e:
            logger.error(f"Error processing {file.name}: {str(e)}")
            st.error(f"Error processing {file.name}")

    def _process_text(self, file) -> Optional[Dict]:
        """Process document text"""
        try:
            # Check file size
            if file.size > MAX_FILE_SIZE:
                raise ValueError(f"File exceeds size limit ({MAX_FILE_SIZE/1024/1024}MB)")

            # Process based on file type
            file_ext = Path(file.name).suffix[1:].lower()
            if file_ext not in self.processors:
                raise ValueError(f"Unsupported file format: {file_ext}")
            
            result = self.processors[file_ext](file)
            if not result or not result.get('content'):
                raise ValueError("No content could be extracted")
            
            return result
            
        except Exception as e:
            logger.error(f"Text processing error: {str(e)}")
            st.error(f"Could not process {file.name}: {str(e)}")
            return None

    def _update_similarities(self, new_file: str, new_doc: Dict) -> None:
        """Update similarity scores for all documents"""
        try:
            calculator = SimilarityCalculator(self.model_manager.similarity_model)
            docs = st.session_state.documents
            
            # Initialize similarities
            if 'similarities' not in new_doc:
                new_doc['similarities'] = {}
            
            # Calculate similarities with existing documents
            for doc_name, doc in docs.items():
                if doc_name != new_file:
                    if 'similarities' not in doc:
                        doc['similarities'] = {}
                    
                    # Calculate similarity
                    scores = calculator.calculate_similarity(
                        new_doc['content'],
                        [doc['content']]
                    )
                    
                    if scores and len(scores) > 0:
                        score = scores[0]
                        # Update both documents
                        new_doc['similarities'][doc_name] = score
                        doc['similarities'][new_file] = score
                        
        except Exception as e:
            logger.error(f"Similarity update error: {str(e)}")
            logger.warning("Could not calculate document similarities")

# processors.py (continued)

class PDFProcessor:
    """Handles PDF document processing"""
    
    @staticmethod
    def process(file) -> Dict[str, str]:
        """Process PDF files with metadata extraction"""
        try:
            pdf_reader = PyPDF2.PdfReader(file)
            text_content = []
            metadata = {}
            
            # Extract PDF metadata
            try:
                info = pdf_reader.metadata
                if info:
                    metadata = {
                        'title': info.get('/Title', ''),
                        'author': info.get('/Author', ''),
                        'subject': info.get('/Subject', ''),
                        'keywords': info.get('/Keywords', ''),
                        'creator': info.get('/Creator', ''),
                        'producer': info.get('/Producer', ''),
                        'creation_date': info.get('/CreationDate', '')
                    }
            except Exception as e:
                logger.warning(f"Metadata extraction failed: {str(e)}")
            
            # Process each page
            for page in pdf_reader.pages:
                page_text = page.extract_text()
                if page_text:
                    cleaned_text = TextProcessor.clean_text(page_text)
                    if cleaned_text:
                        text_content.append(cleaned_text)
            
            # Combine all text
            full_text = '\n'.join(text_content)
            
            if not full_text:
                raise ValueError("No text could be extracted from PDF")
            
            # Calculate statistics
            stats = {
                "type": "pdf",
                "name": file.name,
                "file_size": file.size,
                "num_pages": len(pdf_reader.pages),
                "word_count": len(full_text.split()),
                "char_count": len(full_text)
            }
                
            return {
                "content": full_text,
                "metadata": metadata,
                "stats": stats,
                "processed": True
            }
        except Exception as e:
            logger.error(f"PDF processing error: {str(e)}")
            raise

class DocxProcessor:
    """Handles DOCX document processing"""
    
    @staticmethod
    def process(file) -> Dict[str, str]:
        """Process DOCX files with metadata extraction"""
        try:
            doc = docx.Document(file)
            text_content = []
            metadata = {}
            
            # Extract document properties
            try:
                core_props = doc.core_properties
                metadata = {
                    'title': core_props.title or '',
                    'author': core_props.author or '',
                    'subject': core_props.subject or '',
                    'keywords': core_props.keywords or '',
                    'created': str(core_props.created) if core_props.created else '',
                    'modified': str(core_props.modified) if core_props.modified else '',
                    'last_modified_by': core_props.last_modified_by or ''
                }
            except Exception as e:
                logger.warning(f"Metadata extraction failed: {str(e)}")
            
            # Process paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    formatted_text = [run.text for run in paragraph.runs]
                    clean_text = TextProcessor.clean_text(' '.join(formatted_text))
                    if clean_text:
                        text_content.append(clean_text)
            
            # Process tables
            table_count = 0
            for table in doc.tables:
                table_count += 1
                table_content = []
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = TextProcessor.clean_text(cell.text)
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:
                        table_content.append(' | '.join(row_text))
                if table_content:
                    text_content.append('\n'.join(table_content))
            
            # Combine all content
            full_text = '\n'.join(text_content)
            
            if not full_text:
                raise ValueError("No text could be extracted from DOCX")
            
            # Calculate statistics
            stats = {
                "type": "docx",
                "name": file.name,
                "file_size": file.size,
                "table_count": table_count,
                "paragraph_count": len(doc.paragraphs),
                "word_count": len(full_text.split()),
                "char_count": len(full_text)
            }
                
            return {
                "content": full_text,
                "metadata": metadata,
                "stats": stats,
                "processed": True
            }
        except Exception as e:
            logger.error(f"DOCX processing error: {str(e)}")
            raise

class DocProcessor:
    """Handles DOC document processing"""
    
    @staticmethod
    def process(file) -> Dict[str, str]:
        """Process legacy DOC files"""
        try:
            with tempfile.NamedTemporaryFile(delete=False, suffix='.doc') as temp_file:
                temp_file.write(file.read())
                temp_path = temp_file.name
            
            text = ""
            try:
                # Try antiword first (better formatting)
                text = subprocess.check_output(['antiword', temp_path]).decode('utf-8')
            except (subprocess.SubprocessError, FileNotFoundError):
                try:
                    # Fallback to python-docx
                    doc = docx.Document(temp_path)
                    paragraphs = []
                    for paragraph in doc.paragraphs:
                        if paragraph.text.strip():
                            paragraphs.append(paragraph.text)
                    text = '\n'.join(paragraphs)
                except Exception as e:
                    raise ValueError(f"Could not process DOC file: {str(e)}")
            finally:
                os.unlink(temp_path)
            
            clean_text = TextProcessor.clean_text(text)
            
            if not clean_text:
                raise ValueError("No valid text content in DOC file")
            
            # Calculate statistics
            stats = {
                "type": "doc",
                "name": file.name,
                "file_size": file.size,
                "word_count": len(clean_text.split()),
                "char_count": len(clean_text)
            }
                
            return {
                "content": clean_text,
                "stats": stats,
                "processed": True
            }
            
        except Exception as e:
            logger.error(f"DOC processing error: {str(e)}")
            raise

class TxtProcessor:
    """Handles TXT document processing"""
    
    @staticmethod
    def process(file) -> Dict[str, str]:
        """Process text files with encoding detection"""
        try:
            # Try different encodings
            encodings = ['utf-8', 'latin-1', 'cp1252', 'ascii']
            text = None
            used_encoding = None
            
            for encoding in encodings:
                try:
                    text = file.read().decode(encoding)
                    used_encoding = encoding
                    break
                except UnicodeDecodeError:
                    continue
            
            if text is None:
                raise ValueError("Could not decode text file with supported encodings")
            
            clean_text = TextProcessor.clean_text(text)
            
            if not clean_text:
                raise ValueError("No valid text content in file")
            
            # Calculate statistics
            stats = {
                "type": "txt",
                "name": file.name,
                "file_size": file.size,
                "encoding": used_encoding,
                "line_count": len(text.splitlines()),
                "word_count": len(clean_text.split()),
                "char_count": len(clean_text)
            }
                
            return {
                "content": clean_text,
                "stats": stats,
                "processed": True
            }
        except Exception as e:
            logger.error(f"TXT processing error: {str(e)}")
            raise

class TextProcessor:
    """Utility class for text processing"""
    
    @staticmethod
    def clean_text(text: str) -> str:
        """Clean and normalize text"""
        try:
            if not text:
                return ""
            
            # Remove extra whitespace
            text = re.sub(r'\s+', ' ', text)
            
            # Remove special characters but keep essential punctuation
            text = re.sub(r'[^\w\s.,!?;:()\[\]{}"\'`-]', '', text)
            
            # Remove multiple periods
            text = re.sub(r'\.{2,}', '.', text)
            
            # Fix spacing around punctuation
            text = re.sub(r'\s+([.,!?;:])', r'\1', text)
            
            return text.strip()
            
        except Exception as e:
            logger.error(f"Text cleaning error: {str(e)}")
            return text.strip() if text else ""

    @staticmethod
    def extract_sections(text: str) -> Dict[str, str]:
        """Extract common document sections"""
        try:
            sections = {}
            
            # Common section headers
            patterns = {
                'abstract': r'Abstract[\s\n]+(.+?)(?=\n\s*[A-Z][^a-z]+\n|\Z)',
                'introduction': r'Introduction[\s\n]+(.+?)(?=\n\s*[A-Z][^a-z]+\n|\Z)',
                'methodology': r'(?:Methodology|Methods)[\s\n]+(.+?)(?=\n\s*[A-Z][^a-z]+\n|\Z)',
                'results': r'Results[\s\n]+(.+?)(?=\n\s*[A-Z][^a-z]+\n|\Z)',
                'discussion': r'Discussion[\s\n]+(.+?)(?=\n\s*[A-Z][^a-z]+\n|\Z)',
                'conclusion': r'Conclusion[\s\n]+(.+?)(?=\n\s*[A-Z][^a-z]+\n|\Z)',
                'references': r'References[\s\n]+(.+?)(?=\n\s*[A-Z][^a-z]+\n|\Z)'
            }
            
            for section, pattern in patterns.items():
                match = re.search(pattern, text, re.IGNORECASE | re.DOTALL)
                if match:
                    sections[section] = match.group(1).strip()
            
            return sections
            
        except Exception as e:
            logger.error(f"Section extraction error: {str(e)}")
            return {}